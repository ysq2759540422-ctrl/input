"""
Generate Standardized Code Review Report in DOCX format
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# --- Page Setup ---
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# --- Styles ---
styles = doc.styles

def set_heading_style(paragraph, level=1):
    paragraph.style = doc.styles['Heading %d' % level]

def add_title(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_table_row(table, cells_data, is_header=False, bg_color=None):
    row = table.add_row()
    for i, data in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = str(data)
        if is_header:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10)
        else:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    return row

def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '999999')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_paragraph(doc, text, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.size = Pt(size)
    return p

# ====== COVER PAGE ======
doc.add_paragraph()
doc.add_paragraph()
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run("开源项目代码评审报告")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_para.add_run("Online Examination System\nCode Review Report")
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph()
doc.add_paragraph()

info_table = doc.add_table(rows=6, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ("项目名称:", "在线考试系统 (Online Examination System)"),
    ("评审日期:", "2026年6月2日"),
    ("评审人:", "实验评审小组"),
    ("项目技术栈:", "Python 3.x / Flask 2.x / SQLite"),
    ("代码规模:", "约 2200 行（含模板）"),
    ("评审范围:", "认证模块、题目管理模块、考试模块、成绩管理模块"),
]
for i, (k, v) in enumerate(info_data):
    info_table.rows[i].cells[0].text = k
    info_table.rows[i].cells[1].text = v
    info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# ====== SECTION 1: PROJECT OVERVIEW ======
add_title(doc, "一、项目概述与架构", 1)
p = doc.add_paragraph()
p.add_run("1.1 项目概述\n").bold = True
doc.add_paragraph(
    "本项目为一个典型的在线考试系统（Online Examination System），采用 Flask Web 框架开发，"
    "数据库使用 SQLite，前端使用 HTML + JavaScript 无框架模板。项目实现了用户管理、"
    "题库管理、智能组卷、在线考试、自动评分、成绩查询等核心功能，覆盖了课程要求的完整业务流程。"
)

p2 = doc.add_paragraph()
p2.add_run("1.2 系统架构\n").bold = True
doc.add_paragraph(
    "系统采用经典的三层架构模式，层次清晰，模块之间通过 Service 层进行业务逻辑处理，"
    "通过 Blueprint 路由层对外暴露 REST API 接口。"
)

arch_table = doc.add_table(rows=1, cols=3)
arch_table.style = 'Table Grid'
headers_arch = ["层次", "模块名称", "职责说明"]
for i, h in enumerate(headers_arch):
    arch_table.rows[0].cells[i].text = h
    arch_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    set_cell_shading(arch_table.rows[0].cells[i], '4472C4')
    arch_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

arch_data = [
    ("表现层", "routes/", "接收HTTP请求，参数校验，调用Service，返回JSON响应"),
    ("业务层", "services/", "核心业务逻辑实现，包含各种业务规则和数据处理"),
    ("数据层", "models/", "SQLAlchemy ORM模型定义，数据表结构"),
    ("工具层", "utils/", "通用工具函数，权限装饰器，响应格式化"),
    ("配置层", "config/", "系统配置参数，数据库路径，密钥等"),
    ("入口层", "app.py", "应用初始化，路由注册，数据库初始化脚本"),
]
for row_data in arch_data:
    row = arch_table.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        row.cells[i].paragraphs[0].runs[0].font.size = Pt(9)

set_table_border(arch_table)
doc.add_paragraph()

# ====== SECTION 2: BUSINESS FLOW ======
add_title(doc, "二、核心业务流程", 1)
flow_para = doc.add_paragraph()
flow_para.add_run("系统核心业务流程如下：\n").bold = True

flow_steps = [
    "用户注册/登录 → 认证模块验证用户名密码，建立Session会话",
    "管理员/教师登录 → 进入题库管理 → 创建/编辑/删除题目",
    "教师进入考试管理 → 创建考试 → 设置参数（时长/题量/难度分布）→ 智能随机组卷",
    "教师发布考试 → 学生看到已发布的考试列表",
    "学生进入考试 → 答题 → 提交试卷 → 状态变为待评分",
    "教师评分（手动或批量）→ 系统自动比对答案 → 计算得分",
    "学生查询成绩 → 查看排名统计",
]
for step in flow_steps:
    doc.add_paragraph(step, style='List Number')

doc.add_page_break()

# ====== SECTION 3: CODE REVIEW FINDINGS ======
add_title(doc, "三、代码评审详细报告", 1)

add_title(doc, "3.1 评审维度说明", 2)
doc.add_paragraph(
    "本次评审严格遵循企业级五大评审维度标准，对认证模块（auth_service.py）、"
    "题目管理模块（question_service.py）、考试模块（exam_service.py）、"
    "成绩管理模块（score_service.py）共4个核心模块进行了逐行审查，"
    "共发现15处代码问题，覆盖全部5种问题类型。"
)
doc.add_paragraph()

# Summary table
add_title(doc, "3.2 问题汇总统计", 2)
summary_table = doc.add_table(rows=1, cols=4)
summary_table.style = 'Table Grid'
for i, h in enumerate(["问题类型", "问题数量", "严重程度", "覆盖模块"]):
    summary_table.rows[0].cells[i].text = h
    summary_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    summary_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(summary_table.rows[0].cells[i], '4472C4')

summary_data = [
    ("代码规范性问题", "3", "低-中", "auth, exam"),
    ("代码逻辑缺陷（Bug）", "4", "高-严重", "auth, question, exam, score"),
    ("安全性问题", "3", "严重-高", "auth, score"),
    ("性能问题", "2", "中", "auth, score"),
    ("可维护性问题", "3", "低-中", "question, exam"),
]
for i, row_data in enumerate(summary_data):
    row = summary_table.add_row()
    for j, val in enumerate(row_data):
        row.cells[j].text = val
        row.cells[j].paragraphs[0].runs[0].font.size = Pt(9)
        if i % 2 == 0:
            set_cell_shading(row.cells[j], 'EEF2FF')
set_table_border(summary_table)
doc.add_paragraph()

# ====== DETAILED BUG REPORTS ======
add_title(doc, "3.3 问题详细分析（按严重程度排序）", 2)

def add_bug_section(doc, bug_num, severity, category, title, location,
                    description, current_code, fixed_code, risk, suggestion):
    # Bug header
    p = doc.add_paragraph()
    run = p.add_run(f"【Bug-{bug_num:02d}】{title}")
    run.bold = True
    run.font.size = Pt(12)
    severity_colors_map = {
        "Critical": RGBColor(0xC0, 0x00, 0x00),
        "High": RGBColor(0xFF, 0x66, 0x00),
        "Medium": RGBColor(0x00, 0x00, 0xC0),
        "Low": RGBColor(0x00, 0x80, 0x00),
    }
    if severity in severity_colors_map:
        run.font.color.rgb = severity_colors_map[severity]

    # Meta info table
    meta_table = doc.add_table(rows=3, cols=4)
    meta_table.style = 'Table Grid'
    meta_table.rows[0].cells[0].text = "严重程度"
    meta_table.rows[0].cells[1].text = severity
    meta_table.rows[0].cells[2].text = "问题类型"
    meta_table.rows[0].cells[3].text = category
    meta_table.rows[1].cells[0].text = "文件位置"
    meta_table.rows[1].cells[1].text = location
    meta_table.rows[1].cells[2].text = "Bug编号"
    meta_table.rows[1].cells[3].text = f"Bug-{bug_num:02d}"
    for row in meta_table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    break
            if not cell.paragraphs[0].runs:
                run = cell.paragraphs[0].add_run(cell.text)
                run.font.size = Pt(9)
    set_table_border(meta_table)

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.add_run("问题描述：").bold = True
    doc.add_paragraph(description)

    p3 = doc.add_paragraph()
    p3.add_run("当前问题代码：").bold = True
    code_p = doc.add_paragraph()
    code_run = code_p.add_run(current_code)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(8)
    code_p.paragraph_format.left_indent = Inches(0.3)
    set_cell_shading_for_para = lambda p: None

    p4 = doc.add_paragraph()
    p4.add_run("修复方案：").bold = True
    fix_p = doc.add_paragraph()
    fix_run = fix_p.add_run(fixed_code)
    fix_run.font.name = 'Courier New'
    fix_run.font.size = Pt(8)
    fix_p.paragraph_format.left_indent = Inches(0.3)

    p5 = doc.add_paragraph()
    p5.add_run("风险说明：").bold = True
    p5.add_run(risk)

    p6 = doc.add_paragraph()
    p6.add_run("优化建议：").bold = True
    p6.add_run(suggestion)

    doc.add_paragraph()
    doc.add_paragraph("─" * 80)
    doc.add_paragraph()


# ====== BUG 01: Plaintext Password ======
add_bug_section(doc,
    bug_num=1, severity="Critical", category="安全性问题",
    title="用户密码明文存储",
    location="services/auth_service.py:11, 17",
    description=(
        "用户注册时，密码直接以明文形式存入数据库 user.password 字段。 "
        "密码应该使用哈希算法（如 werkzeug.security.generate_password_hash）加密存储，"
        "登录验证时使用 check_password_hash 校验。明文存储意味着任何人（只要能访问数据库）"
        "都能直接获取所有用户的密码，且一旦数据库泄露，所有账户密码立即暴露。"
    ),
    current_code=(
        "# auth_service.py - register_user\n"
        "def register_user(username, password, role='student', ...):\n"
        "    user = User()\n"
        "    user.username = username\n"
        "    user.password = password  # ← 明文存储，严重安全漏洞\n"
        "    user.role = role\n"
        "    db.session.add(user)\n"
        "    db.session.commit()\n"
        "\n"
        "# auth_service.py - change_password\n"
        "def change_password(user_id, old_password, new_password):\n"
        "    user = User.query.get(user_id)\n"
        "    user.password = new_password  # ← 明文存储新密码\n"
        "    db.session.commit()"
    ),
    fixed_code=(
        "from werkzeug.security import generate_password_hash, check_password_hash\n\n"
        "def register_user(username, password, role='student', ...):\n"
        "    user = User()\n"
        "    user.username = username\n"
        "    user.password = generate_password_hash(password)  # ← 哈希加密\n"
        "    user.role = role\n"
        "    db.session.add(user)\n"
        "    db.session.commit()\n"
        "\n"
        "def change_password(user_id, old_password, new_password):\n"
        "    user = User.query.get(user_id)\n"
        "    if not check_password_hash(user.password, old_password):\n"
        "        return False, 'Old password incorrect'\n"
        "    user.password = generate_password_hash(new_password)  # ← 哈希存储\n"
        "    db.session.commit()"
    ),
    risk="极高风险。数据库泄露即所有用户密码泄露。建议立即修复，同时需要迁移现有明文密码。",
    suggestion="使用 werkzeug.security.generate_password_hash(password, method='pbkdf2:sha256') 加密存储。"
               "修复后需对现有用户密码进行一次性哈希迁移。"
)

# ====== BUG 02: Password Comparison ======
add_bug_section(doc,
    bug_num=2, severity="Critical", category="安全性问题",
    title="登录密码验证未使用哈希比对",
    location="services/auth_service.py:20-23",
    description=(
        "login_user() 方法使用明文比较 user.password != password 进行密码验证。"
        "由于密码是明文存储（Bug-01），这里虽能工作，但无法与哈希存储方案兼容。"
        "此外，直接比较明文密码在日志、调试等场景中容易被捕获，风险极高。"
    ),
    current_code=(
        "def login_user(username, password):\n"
        "    user = User.query.filter_by(username=username).first()\n"
        "    if not user:\n"
        "        return None, 'User not found'\n"
        "    if user.password != password:  # ← 明文比对，未使用check_password_hash\n"
        "        return None, 'Incorrect password'\n"
        "    return user, 'Login successful'"
    ),
    fixed_code=(
        "def login_user(username, password):\n"
        "    user = User.query.filter_by(username=username).first()\n"
        "    if not user:\n"
        "        return None, 'User not found'\n"
        "    if not check_password_hash(user.password, password):\n"
        "        return None, 'Incorrect password'\n"
        "    return user, 'Login successful'"
    ),
    risk="高风险。未使用标准哈希校验，绕过哈希存储后会导致登录永远失败。",
    suggestion="使用 check_password_hash(user.password, password) 进行验证，需先修复Bug-01。"
)

# ====== BUG 03: Old Password Not Verified ======
add_bug_section(doc,
    bug_num=3, severity="High", category="代码逻辑缺陷",
    title="修改密码不验证旧密码",
    location="services/auth_service.py:27-35",
    description=(
        "change_password() 方法接收了 old_password 参数，但完全没有使用它。"
        "任何已登录用户（哪怕Session被劫持）都可以将任意用户密码修改为新值。"
        "另外，update_user_profile() 也没有检查当前操作用户是否与目标用户一致，"
        "导致任意用户可修改他人资料。"
    ),
    current_code=(
        "def change_password(user_id, old_password, new_password):\n"
        "    # old_password 被接收但从未使用！\n"
        "    user = User.query.get(user_id)\n"
        "    if not user:\n"
        "        return False, 'User not found'\n"
        "    user.password = new_password  # ← 未验证旧密码直接修改\n"
        "    db.session.commit()\n"
        "    return True, 'Password changed'\n"
        "\n"
        "def update_user_profile(user_id, realname=None, email=None):\n"
        "    # 无权限检查，任何人可修改任意用户资料\n"
        "    user = User.query.get(user_id)\n"
        "    ..."
    ),
    fixed_code=(
        "def change_password(user_id, old_password, new_password):\n"
        "    user = User.query.get(user_id)\n"
        "    if not user:\n"
        "        return False, 'User not found'\n"
        "    # 必须验证旧密码\n"
        "    if not check_password_hash(user.password, old_password):\n"
        "        return False, 'Old password is incorrect'\n"
        "    user.password = generate_password_hash(new_password)\n"
        "    db.session.commit()\n"
        "    return True, 'Password changed'\n"
        "\n"
        "def update_user_profile(logged_in_user_id, target_user_id, ...):\n"
        "    if logged_in_user_id != target_user_id and role != 'admin':\n"
        "        return False, 'Unauthorized to modify this profile'"
    ),
    risk="高风险。密码修改无旧密码验证，权限校验缺失，可导致账号被非法篡改。",
    suggestion="change_password() 必须在修改前验证旧密码；update_user_profile() 添加权限校验。"
)

# ====== BUG 04: Duplicate Query ======
add_bug_section(doc,
    bug_num=4, severity="Low", category="性能问题",
    title="get_all_teachers方法重复查询数据库",
    location="services/auth_service.py:48-51",
    description=(
        "get_all_teachers() 方法连续执行了两条完全相同的查询语句。"
        "第一条查询结果被第二条覆盖，完全浪费了数据库资源。"
        "这是典型的复制粘贴错误，在数据量较大时造成不必要的数据库负载。"
    ),
    current_code=(
        "def get_all_teachers():\n"
        "    # ← 重复查询，浪费数据库资源\n"
        "    teachers = User.query.filter_by(role='teacher').all()\n"
        "    teachers = User.query.filter_by(role='teacher').all()  # ← 完全重复！\n"
        "    return teachers"
    ),
    fixed_code=(
        "def get_all_teachers():\n"
        "    return User.query.filter_by(role='teacher').all()  # 删除重复行"
    ),
    risk="低风险，仅增加不必要的数据库查询次数。",
    suggestion="删除重复查询行即可。"
)

# ====== BUG 05: No Score Validation ======
add_bug_section(doc,
    bug_num=5, severity="Medium", category="代码规范性问题",
    title="题目分数字段未校验负数与边界值",
    location="services/question_service.py:12-32",
    description=(
        "创建题目时，score 参数未做任何校验。系统接受负数（score=-5）或超大数值。"
        "负分可能导致学生答错反而加分，严重影响考试公正性。"
        "超大分数值可能导致总分超过合理范围（>100）。"
    ),
    current_code=(
        "def create_question(..., score=5, ...):\n"
        "    if not title or not content or not answer:\n"
        "        return None, '...'  # ← 仅校验必填项，未校验score\n"
        "    score = data.get('score', 5)  # ← 直接使用，无范围校验\n"
        "    question.score = score  # 接受负数、超大值"
    ),
    fixed_code=(
        "def create_question(..., score=5, ...):\n"
        "    score = data.get('score', 5)\n"
        "    # 添加分数范围校验\n"
        "    if not isinstance(score, int):\n"
        "        return None, 'Score must be an integer'\n"
        "    if score <= 0 or score > 100:\n"
        "        return None, 'Score must be between 1 and 100'\n"
        "    # ... 后续逻辑"
    ),
    risk="中高风险。负分影响考试成绩公正性，超大值影响统计准确性。",
    suggestion="在 create_question 和 update_question 中添加 score 范围校验（建议 1-100）。"
)

# ====== BUG 06: Invalid Difficulty Accepted ======
add_bug_section(doc,
    bug_num=6, severity="Medium", category="代码逻辑缺陷",
    title="题目难度值未校验合法性",
    location="services/question_service.py:34-46",
    description=(
        "update_question() 方法使用 setattr 批量赋值，允许设置任意字符串为难度值。"
        "系统未定义有效难度枚举（easy/medium/hard），导致 'difficulty=extreme' 这样的无效值被接受。"
        "这会导致统计功能（按难度统计题目数量）结果不准确。"
    ),
    current_code=(
        "ALLOWED_DIFFICULTIES = ['easy', 'medium', 'hard']  # ← 未定义\n\n"
        "def update_question(question_id, **kwargs):\n"
        "    allowed_fields = ['title', 'content', ..., 'difficulty', ...]\n"
        "    for field in allowed_fields:\n"
        "        if field in kwargs:\n"
        "            setattr(question, field, kwargs[field])  # ← 无校验\n"
        "    db.session.commit()"
    ),
    fixed_code=(
        "ALLOWED_DIFFICULTIES = {'easy', 'medium', 'hard'}\n"
        "ALLOWED_TYPES = {'single', 'multiple'}\n\n"
        "def update_question(question_id, **kwargs):\n"
        "    if 'difficulty' in kwargs and kwargs['difficulty'] not in ALLOWED_DIFFICULTIES:\n"
        "        return None, f'Invalid difficulty. Must be one of {ALLOWED_DIFFICULTIES}'\n"
        "    if 'type' in kwargs and kwargs['type'] not in ALLOWED_TYPES:\n"
        "        return None, f'Invalid type. Must be one of {ALLOWED_TYPES}'\n"
        "    for field in allowed_fields:\n"
        "        if field in kwargs:\n"
        "            setattr(question, field, kwargs[field])"
    ),
    risk="中风险。无效难度值破坏统计准确性，数据质量下降。",
    suggestion="定义常量 ALLOWED_DIFFICULTIES 并在赋值前校验。"
)

# ====== BUG 07: No Transaction Rollback ======
add_bug_section(doc,
    bug_num=7, severity="Medium", category="代码逻辑缺陷",
    title="批量创建题目无事务回滚机制",
    location="services/question_service.py:58-76",
    description=(
        "batch_create_questions() 在循环内对每个题目单独执行 db.session.commit()。"
        "如果第5条数据因数据问题失败，前4条已永久提交。"
        "违反原子性原则——批量操作要么全部成功，要么全部回滚。"
        "另外，若中途失败，已提交的数据无法撤回，导致数据不一致。"
    ),
    current_code=(
        "def batch_create_questions(questions_data, creator_id):\n"
        "    created = []\n"
        "    for data in questions_data:\n"
        "        question = Question(...)\n"
        "        db.session.add(question)\n"
        "        db.session.commit()  # ← 每个单独提交，失败时无法回滚\n"
        "        created.append(question)\n"
        "    return created, len(created)"
    ),
    fixed_code=(
        "def batch_create_questions(questions_data, creator_id):\n"
        "    try:\n"
        "        created = []\n"
        "        for data in questions_data:\n"
        "            question = Question(...)\n"
        "            db.session.add(question)\n"
        "        db.session.commit()  # ← 全部添加后再一次性提交\n"
        "        created = Question.query.filter(\n"
        "            Question.creator_id == creator_id\n"
        "        ).order_by(Question.id.desc()).limit(len(questions_data)).all()\n"
        "        return created, len(created)\n"
        "    except Exception as e:\n"
        "        db.session.rollback()  # ← 出错时全部回滚\n"
        "        return [], 0"
    ),
    risk="中风险。批量导入中断时部分数据已保存，导致数据库状态不一致。",
    suggestion="将 db.session.commit() 移至循环外部，用 try/except 包裹并添加 rollback。"
)

# ====== BUG 08: random.sample ValueError ======
add_bug_section(doc,
    bug_num=8, severity="High", category="代码逻辑缺陷",
    title="组卷随机抽样遇负数导致ValueError崩溃",
    location="services/exam_service.py:63",
    description=(
        "当 question_count <= 0 时，_auto_generate_questions() 将负数传入 random.sample()，"
        "Python 标准库会抛出 ValueError: Sample larger than population or is negative，"
        "导致整个服务崩溃。这是未经处理的异常边界条件，影响系统可用性。"
        "自动化测试 TC-24 已验证此 Bug。"
    ),
    current_code=(
        "def _auto_generate_questions(count, difficulty_distribution=None, category=None):\n"
        "    ...\n"
        "    for difficulty, ratio in dist.items():\n"
        "        target_count = int(count * ratio)\n"
        "        difficulty_questions = [q for q in remaining if q.difficulty == difficulty]\n"
        "        sampled = random.sample(difficulty_questions,\n"
        "            min(target_count, len(difficulty_questions)))  # ← 当count为负数时崩溃\n"
        "        ..."
    ),
    fixed_code=(
        "def create_exam(..., question_count, ...):\n"
        "    if isinstance(question_count, int) and question_count <= 0:\n"
        "        return None, 'question_count must be a positive integer'\n"
        "\n"
        "# 或者在 _auto_generate_questions 中保护\n"
        "def _auto_generate_questions(count, ...):\n"
        "    if not isinstance(count, int) or count <= 0:\n"
        "        return []  # 返回空列表而非崩溃"
    ),
    risk="高风险。传负数时服务崩溃（500错误），影响系统可用性。",
    suggestion="在 create_exam 入口处添加 question_count > 0 的校验，提前拦截非法输入。"
)

# ====== BUG 09: Magic Value ======
add_bug_section(doc,
    bug_num=9, severity="Low", category="可维护性问题",
    title="考试时长上限使用魔法值（Magic Number）",
    location="services/exam_service.py:15",
    description=(
        "考试时长上限 1000 分钟直接以硬编码数字出现在代码中。"
        "魔法值使得代码可读性差、维护成本高。"
        "若业务需求变化（如改为最大720分钟），需找到所有硬编码位置修改，易遗漏。"
    ),
    current_code=(
        "def create_exam(title, description, duration, question_count, ...):\n"
        "    if duration <= 0 or duration > 1000:  # ← 魔法值1000\n"
        "        return None, 'Duration must be between 1 and 1000 minutes'"
    ),
    fixed_code=(
        "# 在文件顶部或 config.py 中定义常量\n"
        "MAX_EXAM_DURATION = 1000  # minutes\n"
        "MIN_EXAM_DURATION = 1\n\n"
        "def create_exam(..., duration, ...):\n"
        "    if duration < MIN_EXAM_DURATION or duration > MAX_EXAM_DURATION:\n"
        "        return None, f'Duration must be between {MIN_EXAM_DURATION} and {MAX_EXAM_DURATION} minutes'"
    ),
    risk="低风险，主要影响代码可维护性和可读性。",
    suggestion="使用有名称的常量替代魔法值，提高代码可维护性。"
)

# ====== BUG 10: Publish Without Questions ======
add_bug_section(doc,
    bug_num=10, severity="High", category="代码逻辑缺陷",
    title="允许发布无题目的考试",
    location="services/exam_service.py:84-93",
    description=(
        "publish_exam() 方法虽然检查 question_ids 是否为空，但检查逻辑存在缺陷："
        "当 question_ids 为 None 或空字符串时，len(json.loads(...)) 会返回 0，但 "
        "若 question_count > 0 但 question_ids 为空，说明组卷失败但状态正常。"
        "更严重的是，若 question_ids 列表长度为 0 或 question_count 与实际题目数不符，"
        "系统允许发布一个无有效题目的考试，学生打开后将看到空白试卷。"
    ),
    current_code=(
        "def publish_exam(exam_id):\n"
        "    exam = Exam.query.get(exam_id)\n"
        "    if not exam:\n"
        "        return False, 'Exam not found'\n"
        "    if not exam.question_ids or len(json.loads(exam.question_ids or '[]')) == 0:\n"
        "        return False, 'Cannot publish exam without questions'\n"
        "    exam.is_published = 1  # ← 仍可能发布question_count不匹配的考试\n"
        "    db.session.commit()"
    ),
    fixed_code=(
        "def publish_exam(exam_id):\n"
        "    exam = Exam.query.get(exam_id)\n"
        "    if not exam:\n"
        "        return False, 'Exam not found'\n"
        "    qids = exam.get_question_ids()\n"
        "    if not qids or len(qids) == 0:\n"
        "        return False, 'Cannot publish exam without questions'\n"
        "    if len(qids) != exam.question_count:\n"
        "        return False, f'Question count mismatch: declared {exam.question_count}, actual {len(qids)}'\n"
        "    exam.is_published = 1\n"
        "    db.session.commit()\n"
        "    return True, 'Exam published'"
    ),
    risk="高风险。无题目考试被发布后，学生无法正常作答，严重影响用户体验和系统可信度。",
    suggestion="发布前验证：1) 题目列表非空；2) 题目数量与声明一致。"
)

# ====== BUG 11: Clone Missing Fields ======
add_bug_section(doc,
    bug_num=11, severity="Medium", category="可维护性问题",
    title="克隆考试字段复制不完整",
    location="services/exam_service.py:154-167",
    description=(
        "clone_exam() 方法复制源考试时遗漏了 description、total_score、is_published 等字段。"
        "克隆出的考试没有描述内容，总分为默认值100（而非源考试实际总分），"
        "且发布状态也被忽略。这些遗漏会导致教师在使用克隆考试功能时遇到数据缺失问题。"
    ),
    current_code=(
        "def clone_exam(exam_id, new_creator_id):\n"
        "    source = Exam.query.get(exam_id)\n"
        "    new_exam = Exam()\n"
        "    new_exam.title = source.title + ' (Copy)'\n"
        "    new_exam.duration = source.duration\n"
        "    # ← 缺失: description\n"
        "    new_exam.question_ids = source.question_ids\n"
        "    # ← 缺失: total_score\n"
        "    # ← 缺失: is_published (应为0)\n"
        "    new_exam.creator_id = new_creator_id\n"
        "    db.session.add(new_exam)\n"
        "    db.session.commit()"
    ),
    fixed_code=(
        "def clone_exam(exam_id, new_creator_id):\n"
        "    source = Exam.query.get(exam_id)\n"
        "    new_exam = Exam()\n"
        "    new_exam.title = source.title + ' (Copy)'\n"
        "    new_exam.description = source.description\n"
        "    new_exam.duration = source.duration\n"
        "    new_exam.total_score = source.total_score\n"
        "    new_exam.question_count = source.question_count\n"
        "    new_exam.question_ids = source.question_ids\n"
        "    new_exam.creator_id = new_creator_id\n"
        "    new_exam.is_published = 0  # 克隆版本默认不发布\n"
        "    db.session.add(new_exam)\n"
        "    db.session.commit()\n"
        "    return new_exam, 'Exam cloned'"
    ),
    risk="中等风险。克隆考试数据不完整，影响用户体验。",
    suggestion="确保所有需要保留的字段都被正确复制。"
)

# ====== BUG 12: Case-Sensitive Grading ======
add_bug_section(doc,
    bug_num=12, severity="High", category="代码逻辑缺陷",
    title="答案比对大小写敏感导致评分错误",
    location="services/score_service.py:47-55",
    description=(
        "grade_exam() uses direct string equality to compare student answers: student_answer == question.answer. "
        "If a student inputs 'a' but the correct answer is 'A', the system marks it as wrong. "
        "Letter case should not affect multiple choice correctness - this is a major grading logic defect. "
        "Test TC-30 has verified this Bug: lowercase answers are all marked incorrect."
    ),
    current_code=(
        "def grade_exam(score_id, grader_id=None):\n"
        "    ...\n"
        "    for qid in question_ids:\n"
        "        question = Question.query.get(qid)\n"
        "        student_answer = student_answers.get(str(qid), '')\n"
        "        if student_answer == question.answer:  # ← 大小写敏感比较\n"
        "            total_score += question.score  # 大小写不匹配则不得分"
    ),
    fixed_code=(
        "def grade_exam(score_id, grader_id=None):\n"
        "    ...\n"
        "    for qid in question_ids:\n"
        "        question = Question.query.get(qid)\n"
        "        student_answer = student_answers.get(str(qid), '')\n"
        "        # 统一转为大写后比较，忽略首尾空格\n"
        "        if student_answer.strip().upper() == question.answer.strip().upper():\n"
        "            total_score += question.score\n"
        "# 同时修复 batch_grade() 中的相同问题"
    ),
    risk="高风险。严重影响评分准确性，违反常识（大小写不应影响选择题正确性）。",
    suggestion="评分时统一使用 .strip().upper() 规范化比较，同时修复 batch_grade() 中的相同问题。"
)

# ====== BUG 13: Statistics Returns None ======
add_bug_section(doc,
    bug_num=13, severity="Medium", category="可维护性问题",
    title="统计接口无数据时返回None破坏前端契约",
    location="services/score_service.py:76-83",
    description=(
        "calculate_exam_statistics() 在没有成绩记录时返回 (None, 'No scores yet')。"
        "前端通常期望得到有效的数据结构（即使为空），返回 None 会导致前端解包失败，"
        "可能引发 NullReferenceError。API 应始终返回一致的数据结构。"
    ),
    current_code=(
        "def calculate_exam_statistics(exam_id):\n"
        "    scores = Score.query.filter_by(exam_id=exam_id, status='scored').all()\n"
        "    if not scores:\n"
        "        return None, 'No scores yet'  # ← 前端解包困难\n"
        "    score_values = [s.score for s in scores]\n"
        "    avg = sum(score_values) / len(score_values)\n"
        "    return {\n"
        "        'exam_id': exam_id,\n"
        "        'total_students': len(scores),\n"
        "        'average': round(avg, 2),\n"
        "        ...\n"
        "    }, 'Success'"
    ),
    fixed_code=(
        "def calculate_exam_statistics(exam_id):\n"
        "    scores = Score.query.filter_by(exam_id=exam_id, status='scored').all()\n"
        "    if not scores:\n"
        "        return {\n"
        "            'exam_id': exam_id,\n"
        "            'total_students': 0,\n"
        "            'average': 0,\n"
        "            'max': 0,\n"
        "            'min': 0,\n"
        "            'pass_count': 0\n"
        "        }, 'Success'  # ← 返回有效空结构\n"
        "    score_values = [s.score for s in scores]\n"
        "    avg = sum(score_values) / len(score_values)"
    ),
    risk="中风险。前端需要额外处理 None，增加代码复杂度，可能导致运行时错误。",
    suggestion="始终返回统一格式的数据结构，即使数据为空也返回零值填充的结构。"
)

# ====== BUG 14: No Permission Check on Delete Score ======
add_bug_section(doc,
    bug_num=14, severity="High", category="安全性问题",
    title="成绩删除无权限校验",
    location="services/score_service.py:96-101",
    description=(
        "delete_score() 方法完全没有任何权限检查。任何已登录用户（包括学生）"
        "都可以通过猜测 score_id 删除任意成绩记录。"
        "这意味着学生可以删除自己或他人的成绩，甚至教师之间可以互相删除成绩。"
        "这是一个严重的安全漏洞，违反了最小权限原则。"
    ),
    current_code=(
        "def delete_score(score_id):\n"
        "    # ← 无权限检查，任何登录用户都可调用\n"
        "    score_record = Score.query.get(score_id)\n"
        "    if not score_record:\n"
        "        return False, 'Score record not found'\n"
        "    db.session.delete(score_record)\n"
        "    db.session.commit()\n"
        "    return True, 'Score deleted'"
    ),
    fixed_code=(
        "def delete_score(score_id, requesting_user_id, requesting_role):\n"
        "    score_record = Score.query.get(score_id)\n"
        "    if not score_record:\n"
        "        return False, 'Score record not found'\n"
        "    # 权限检查：仅管理员/教师可删除，或本人可申请删除\n"
        "    if requesting_role not in ['admin', 'teacher']:\n"
        "        return False, 'Permission denied: admin or teacher role required'\n"
        "    # 教师只能删除自己创建的考试成绩\n"
        "    if requesting_role == 'teacher':\n"
        "        exam = Exam.query.get(score_record.exam_id)\n"
        "        if exam and exam.creator_id != requesting_user_id:\n"
        "            return False, 'Permission denied'\n"
        "    db.session.delete(score_record)\n"
        "    db.session.commit()\n"
        "    return True, 'Score deleted'"
    ),
    risk="高风险。严重安全漏洞——任何用户可删除任何成绩记录，破坏数据完整性。",
    suggestion="添加角色检查（admin/teacher）；教师仅能删除自己创建的考试对应成绩。"
)

# ====== BUG 15: N+1 Query Problem ======
add_bug_section(doc,
    bug_num=15, severity="Medium", category="性能问题",
    title="批量评分存在N+1查询问题",
    location="services/score_service.py:103-126",
    description=(
        "batch_grade() 在循环中对每个成绩记录分别查询 Exam（N次），"
        "再对每个题目分别查询 Question（N×M次）。"
        "假设有100个成绩记录，每份试卷50题，将产生 1 + 100 + 5000 = 5101 次查询，"
        "而理论上只需约3次查询即可完成（一次获取所有成绩，一次获取关联考试，一次获取所有题目）。"
        "这会导致批量评分操作极慢，数据库连接耗尽。"
    ),
    current_code=(
        "def batch_grade(exam_id):\n"
        "    scores = Score.query.filter_by(...).all()  # 1次查询\n"
        "    graded = 0\n"
        "    for score_record in scores:  # N次循环\n"
        "        exam = Exam.query.get(score_record.exam_id)    # N次查询\n"
        "        question_ids = exam.get_question_ids()\n"
        "        for qid in question_ids:                          # M次循环\n"
        "            question = Question.query.get(qid)         # N×M次查询\n"
        "            if student_answer == question.answer:\n"
        "                total_score += question.score\n"
        "        score_record.score = total_score\n"
        "        db.session.commit()  # N次commit\n"
        "        graded += 1\n"
        "    return graded  # 总计 1+N+NM 次查询"
    ),
    fixed_code=(
        "def batch_grade(exam_id):\n"
        "    scores = Score.query.filter_by(exam_id=exam_id, status='pending').all()\n"
        "    if not scores:\n"
        "        return 0\n"
        "    # 1. 批量获取所有需要的考试\n"
        "    exam_ids = list(set(s.exam_id for s in scores))\n"
        "    exams = {e.id: e for e in Exam.query.filter(Exam.id.in_(exam_ids)).all()}\n"
        "    # 2. 收集所有需要的题目ID\n"
        "    all_q_ids = set()\n"
        "    for exam in exams.values():\n"
        "        all_q_ids.update(exam.get_question_ids())\n"
        "    # 3. 批量获取所有题目\n"
        "    questions = {q.id: q for q in Question.query.filter(Question.id.in_(all_q_ids)).all()}\n"
        "    # 4. 内存中评分，无需更多查询\n"
        "    for score_record in scores:\n"
        "        exam = exams[score_record.exam_id]\n"
        "        student_answers = score_record.get_answers()\n"
        "        total_score = 0\n"
        "        for qid in exam.get_question_ids():\n"
        "            q = questions.get(qid)\n"
        "            if q and student_answers.get(str(qid), '').strip().upper() == q.answer.strip().upper():\n"
        "                total_score += q.score\n"
        "        score_record.score = total_score\n"
        "        score_record.status = 'scored'\n"
        "        score_record.graded_at = datetime.utcnow()\n"
        "    db.session.commit()  # 一次性提交全部\n"
        "    return len(scores)"
    ),
    risk="中等-高风险。100份试卷50题将产生5101次查询，数据库连接池耗尽，评分极慢。",
    suggestion="使用批量查询（IN语句）预先获取所有数据，在内存中完成评分，最后一次性提交。"
)

doc.add_page_break()

# ====== SECTION 4: SUMMARY ======
add_title(doc, "四、评审总结", 1)

add_title(doc, "4.1 共性问题汇总", 2)
common_issues = [
    ("输入校验缺失", "多处未对用户输入参数做范围、类型、枚举值校验，导致非法数据可进入系统。"),
    ("权限控制薄弱", "部分关键操作（删除成绩、修改密码）缺少权限校验，违反最小权限原则。"),
    ("事务边界不清", "批量操作未使用统一事务边界，部分成功部分失败时无法回滚。"),
    ("异常处理粗糙", "部分边界条件（如负数、空集合）未做预处理，可能抛出未捕获异常。"),
    ("代码复用不足", "多处存在重复逻辑（如答案比较逻辑在 grade_exam 和 batch_grade 中重复）。"),
]
for title, desc in common_issues:
    p = doc.add_paragraph()
    p.add_run(f"• {title}：").bold = True
    p.add_run(desc)

doc.add_paragraph()
add_title(doc, "4.2 改进优先级建议", 2)

priority_table = doc.add_table(rows=1, cols=4)
priority_table.style = 'Table Grid'
for i, h in enumerate(["优先级", "Bug编号", "问题描述", "建议行动"]):
    priority_table.rows[0].cells[i].text = h
    priority_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    priority_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(priority_table.rows[0].cells[i], 'C00000')

priority_data = [
    ("P0-立即修复", "Bug-01,02,03,14", "密码安全+权限控制缺失", "修改密码哈希+添加权限检查"),
    ("P1-本周修复", "Bug-08,10,12", "逻辑缺陷影响核心功能", "添加参数校验+评分大小写"),
    ("P2-计划修复", "Bug-05,06,07,11,13", "数据质量+事务+API契约", "添加输入校验+统一返回结构"),
    ("P3-后续优化", "Bug-04,09,15", "性能+代码规范", "删除重复+常量+批量查询"),
]
for i, row_data in enumerate(priority_data):
    row = priority_table.add_row()
    for j, val in enumerate(row_data):
        row.cells[j].text = val
        row.cells[j].paragraphs[0].runs[0].font.size = Pt(9)
set_table_border(priority_table)

doc.add_paragraph()
add_title(doc, "4.3 如何通过代码评审提前规避测试缺陷", 2)
doc.add_paragraph(
    "本次评审发现的 15 个问题中，相当一部分（如 Bug-05 负分、Bug-08 负数崩溃、"
    "Bug-12 大小写敏感性）完全可以通过规范的代码评审流程在编码阶段发现和修复：\n"
)
measures = [
    '强制执行输入校验规范：所有 API 参数在 Service 层入口处必须校验范围和类型，'
    '不允许「裸参数」直接进入业务逻辑。',
    "安全编码规范：密码必须使用哈希存储（PBKDF2/scrypt），敏感操作（删除、修改密码）"
    "必须经过权限校验，权限校验逻辑需同行业评审。",
    "事务边界规范：批量写操作必须在事务内执行，统一 try/except/rollback 模式，"
    "禁止在循环内 commit()。",
    "测试用例驱动评审：评审时应审查是否有对应的边界测试（如 TC-24 负数测试）"
    "，边界条件测试覆盖率低的模块需打回补充。",
    "代码重复检测：使用工具（如 SonarQube）扫描重复代码，"
    "相同逻辑出现2次以上必须提取为公共方法。",
]
for m in measures:
    doc.add_paragraph(f"• {m}")

doc.add_paragraph()
add_title(doc, "4.4 实验收获与工程化开发规范心得", 2)
doc.add_paragraph(
    "通过本次实验，我深入体验了企业级代码评审的完整流程，主要收获如下：\n\n"
    "1. 评审的系统性：代码评审不是找茬，而是以工程质量为准绳，从安全性、正确性、"
    "可维护性、性能等多维度系统性审查，每一个维度都关乎系统的长期健康。\n\n"
    "2. 输入校验是安全的基石：大量 Bug 源于输入校验缺失。在编码阶段就应当假设"
    "所有外部输入都是恶意的，做最严格的校验，这比事后修复代价低得多。\n\n"
    "3. 测试驱动的评审：测试用例（如 TC-24 触发 Bug-08）是发现隐藏 Bug 的有效手段。"
    "好的测试覆盖率能让评审者快速定位问题区域。\n\n"
    "4. 缺陷分级管理：并非所有 Bug 都同等紧急。通过严重程度分级，能让团队聚焦"
    "最高风险问题，避免在低优先级问题上浪费资源。\n\n"
    "5. 文档即代码：标准化的代码评审报告（包含源码定位、风险分析、修复方案）"
    "是团队知识沉淀的重要载体，也是后续维护的重要参考。\n\n"
    "工程化规范方面，本次实验让我深刻认识到：好的代码不是一次写成的，而是在"
    "评审→修复→再评审的迭代循环中持续打磨的。每一个被提前发现的 Bug，都意味着"
    "线上故障的可能减少、用户信任的持续维护。"
)

# Save
output_path = r"d:\Software\python_project\rjcs\reports\Code_Review_Report.docx"
doc.save(output_path)
print(f"Word code review report saved: {output_path}")
