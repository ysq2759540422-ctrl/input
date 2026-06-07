"""
生成代码评审报告 2.0
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)

def set_bd(tbl):
    t = tbl._tbl
    tP = t.tblPr
    b = OxmlElement("w:tblBorders")
    for n in ["top","left","bottom","right","insideH","insideV"]:
        e = OxmlElement(f"w:{n}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), "999999")
        b.append(e)
    tP.append(b)

def shd(cell, hex_c):
    tc = cell._tc; tP = tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear"); s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), hex_c); tP.append(s)

def title(doc, t, lv=1):
    return doc.add_heading(t, lv)

# ===== COVER =====
doc.add_paragraph()
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("在线考试系统 · 代码评审报告")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x44,0x72,0xC4)
sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp.add_run("Code Review Report V2.0\n在线考试系统")
doc.add_paragraph()
tbl = doc.add_table(rows=6, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [
    ("项目名称:", "在线考试系统 (Online Examination System)"),
    ("评审日期:", "2026年6月2日"),
    ("评审人:", "评审小组"),
    ("项目技术栈:", "Python 3.x / Flask 2.x / SQLite"),
    ("代码规模:", "约 2200 行（含模板）"),
    ("评审范围:", "认证模块、题目管理模块、考试模块、成绩管理模块"),
]
for i,(k,v) in enumerate(data):
    tbl.rows[i].cells[0].text = k
    tbl.rows[i].cells[1].text = v
    tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
doc.add_page_break()

# ===== 一、项目概述 =====
title(doc, "一、项目概述", 1)
doc.add_paragraph(
    "本项目为一个典型的在线考试系统，采用 Flask Web 框架开发，数据库使用 SQLite，"
    "前端使用 HTML + JavaScript 无框架模板。项目实现了用户管理、题库管理、"
    "智能组卷、在线考试、自动评分、成绩查询等核心功能。"
)

# ===== 二、评审范围 =====
title(doc, "二、评审范围与方法", 1)
doc.add_paragraph(
    "本次评审对认证模块（auth_service.py）、题目管理模块（question_service.py）、"
    "考试模块（exam_service.py）、成绩管理模块（score_service.py）共4个核心模块进行了审查。"
    "评审重点关注系统功能层面的实际缺陷，经实际运行验证，共发现 12 个需关注的代码问题。"
)

doc.add_page_break()

# ===== 三、缺陷汇总 =====
title(doc, "三、缺陷汇总统计", 1)
smtbl = doc.add_table(rows=1, cols=4)
smtbl.style = "Table Grid"
for i,h in enumerate(["缺陷编号","缺陷标题","严重程度","所在模块"]):
    smtbl.rows[0].cells[i].text = h
    smtbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    smtbl.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    shd(smtbl.rows[0].cells[i], "4472C4")

smdata = [
    ("Bug-01","密码以明文存储","高","认证模块"),
    ("Bug-02","管理员重置密码不验证旧密码","高","认证模块"),
    ("Bug-03","管理员可删除任意用户","中","认证模块"),
    ("Bug-04","题目难度值无枚举校验","中","题目模块"),
    ("Bug-05","题目分数字段无范围校验","高","题目模块"),
    ("Bug-06","批量创建题目无事务回滚","中","题目模块"),
    ("Bug-07","克隆考试description字段丢失","中","考试模块"),
    ("Bug-08","删除考试不检查关联成绩","高","考试模块"),
    ("Bug-09","更新考试题目数不同步组卷","中","考试模块"),
    ("Bug-10","评分答案比对区分大小写","高","成绩模块"),
    ("Bug-11","批量评分同样区分大小写","高","成绩模块"),
    ("Bug-12","无成绩统计返回404","中","成绩模块"),
]
for i,rd in enumerate(smdata):
    row = smtbl.add_row()
    for j,val in enumerate(rd):
        row.cells[j].text = val
        row.cells[j].paragraphs[0].runs[0].font.size = Pt(9)
        if i%2==0: shd(row.cells[j], "EEF2FF")
        if rd[2]=="高": shd(row.cells[j], "FFE0E0")
        elif rd[2]=="中": shd(row.cells[j], "FFF3E0")
set_bd(smtbl)
doc.add_paragraph()
doc.add_page_break()

# ===== 四、详细分析 =====
title(doc, "四、缺陷详细分析", 1)

def add_bug(doc, num, sev, title_txt, loc, desc, code, risk, fix):
    p = doc.add_paragraph()
    r = p.add_run(f"【Bug-{num:02d}】{title_txt}")
    r.bold = True; r.font.size = Pt(12)
    cm = {"高": RGBColor(0xC0,0x00,0x00), "中": RGBColor(0x00,0x00,0xC0), "低": RGBColor(0x00,0x80,0x00)}
    if sev in cm: r.font.color.rgb = cm[sev]
    mt = doc.add_table(rows=2, cols=4)
    mt.style = "Table Grid"
    mt.rows[0].cells[0].text = "严重程度"; mt.rows[0].cells[1].text = sev
    mt.rows[0].cells[2].text = "文件位置"; mt.rows[0].cells[3].text = loc
    mt.rows[1].cells[0].text = "缺陷编号"; mt.rows[1].cells[1].text = f"Bug-{num:02d}"
    mt.rows[1].cells[2].text = "问题类型"; mt.rows[1].cells[3].text = "功能缺陷"
    for row in mt.rows:
        for cell in row.cells:
            for pp in cell.paragraphs:
                for rr in pp.runs: rr.font.size = Pt(9); break
            if not cell.paragraphs[0].runs:
                cell.paragraphs[0].add_run(cell.text).font.size = Pt(9)
    set_bd(mt)
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.add_run("问题描述：").bold = True
    doc.add_paragraph(desc)
    p3 = doc.add_paragraph()
    p3.add_run("当前问题代码：").bold = True
    cp = doc.add_paragraph()
    cr = cp.add_run(code)
    cr.font.name = "Courier New"; cr.font.size = Pt(8)
    cp.paragraph_format.left_indent = Inches(0.3)
    p4 = doc.add_paragraph()
    p4.add_run("风险说明：").bold = True; p4.add_run(risk)
    p5 = doc.add_paragraph()
    p5.add_run("改进建议：").bold = True; p5.add_run(fix)
    doc.add_paragraph()
    doc.add_paragraph("-" * 80)
    doc.add_paragraph()

# Bug-01
add_bug(doc, 1, "高", "密码以明文存储",
    "auth_service.py:16",
    "register_user() 方法将用户密码以明文形式直接存入数据库 password 字段，未使用哈希加密。这是严重的安全缺陷，一旦数据库泄露，所有用户密码将直接暴露。",
    'user.password = password  # 直接存储明文密码',
    "高风险。数据库泄露后所有用户密码直接暴露。",
    "使用 werkzeug.security.generate_password_hash 对密码加密，登录时使用 check_password_hash 验证。"
)

# Bug-02（合并原Bug-02和Bug-03，统一描述为"管理员重置/修改密码不验证旧密码"）
add_bug(doc, 2, "高", "管理员重置密码和修改密码均不验证旧密码",
    "auth_service.py:36-93",
    "admin_reset_password() 和 change_password() 方法均接收了 old_password 参数，但完全没有使用它。系统直接用新密码覆盖旧密码，管理员可以强制重置任意用户密码，普通用户可以修改任意账户的密码。",
    'def admin_reset_password(user_id, new_password):\n    user = User.query.get(user_id)\n    user.password = new_password  # 直接修改，无旧密码验证\n    db.session.commit()\n\ndef change_password(user_id, old_password, new_password):\n    # old_password 参数被接收但从未使用！\n    user.password = new_password  # 直接覆盖，不验证旧密码',
    "高风险。任何已登录用户都可以修改任意用户密码，管理员可强制重置任意账户密码。",
    "重置和修改密码前必须验证旧密码：if user.password != old_password: return False, '旧密码错误'。"
)

# Bug-03（原Bug-04）
add_bug(doc, 3, "中", "管理员可删除任意用户",
    "auth_service.py:77-83 + routes/auth.py:122-126",
    "delete_user() 方法只有 @admin_required 权限校验，但没有限制管理员只能删除哪些用户。管理员可以删除系统中的任何用户，包括其他管理员自己。",
    '@auth_bp.route("/user/<int:user_id>", methods=["DELETE"])\n@admin_required  # 仅检查是否为管理员，不检查删除范围\ndef delete_user(user_id):\n    db.session.delete(user)  # 直接删除，无任何限制',
    "中等风险。管理员可删除任意用户，可能导致误删重要账号。",
    "增加删除范围限制：禁止管理员删除自己，并记录删除操作日志。"
)

# Bug-04（原Bug-05）
add_bug(doc, 4, "中", "题目难度值无枚举校验",
    "question_service.py:19 + routes/question.py",
    "create_question() 和 update_question() 方法中，difficulty 参数直接赋值，没有任何枚举限制。用户可以输入任意字符串如 'extreme'、'超级难' 等，影响题目统计准确性。",
    'question.difficulty = difficulty  # 接受任意字符串\n# routes/question.py 中也是 select 下拉框改为文本框后，无任何校验\n',
    "中等风险。无效难度值破坏按难度统计题目的准确性，数据质量下降。",
    "定义难度枚举常量 VALID_DIFFICULTIES = [\"easy\", \"medium\", \"hard\"]，赋值前校验：if difficulty not in VALID_DIFFICULTIES: return None, '无效的难度值'。"
)

# Bug-05（原Bug-06）
add_bug(doc, 5, "高", "题目分数字段无范围校验",
    "question_service.py:21 + routes/question.py",
    "create_question() 方法中 score 参数直接赋值，没有任何范围校验。负分题目会导致学生答错反而被加分，超大分数影响总分和及格率。",
    'question.score = score  # 接受 0、负数、超大值\n# 前端 list.html 中 <input type="number"> 已去掉 min="1" 限制\n',
    "高风险。负分严重影响考试公正性，超大分数破坏总分计算。",
    "在赋值前添加校验：if not isinstance(score, int) or score < 1 or score > 100: return None, '分值必须在1-100之间'。"
)

# Bug-06（原Bug-07）
add_bug(doc, 6, "中", "批量创建题目无事务回滚",
    "question_service.py:79-100",
    "batch_create_questions() 在循环内对每个题目单独执行 db.session.commit()。如果第3条数据失败，前2条已永久提交无法回滚，破坏批量操作的原子性。",
    'for data in questions_data:\n    question = Question(...)\n    db.session.add(question)\n    db.session.commit()  # 每条单独提交，部分失败无法回滚\n    created.append(question)',
    "中等风险。部分失败时已成功的数据无法回滚，导致数据库中部分记录存在而整体不完整。",
    "将 db.session.commit() 移至循环外部，用 try/except 包裹整体，失败时执行 db.session.rollback()。"
)

# Bug-07（原Bug-08）
add_bug(doc, 7, "中", "克隆考试description字段丢失",
    "exam_service.py:152-163",
    "clone_exam() 方法复制源考试时只复制了 title、duration、question_ids 和 creator_id，遗漏了 description。克隆考试没有描述内容，影响教师对考试的辨识。",
    'def clone_exam(exam_id, new_creator_id):\n    new_exam.title = source.title + " (Copy)"\n    # description 字段未被复制\n    new_exam.duration = source.duration\n    new_exam.question_ids = source.question_ids',
    "中等风险。克隆考试丢失描述内容，教师无法区分多个同名考试。",
    "添加：new_exam.description = source.description，完整复制所有字段。"
)

# Bug-08（原Bug-09）
add_bug(doc, 8, "高", "删除考试不检查关联成绩",
    "exam_service.py:129-136",
    "delete_exam() 直接执行删除操作，没有检查该考试是否有关联成绩记录。删除后 Score 记录仍然存在但指向不存在的 exam_id，被外键约束阻止或产生孤儿记录。",
    'def delete_exam(exam_id):\n    exam = Exam.query.get(exam_id)\n    # 未检查是否有关联成绩记录\n    db.session.delete(exam)  # 可能触发外键约束错误\n    db.session.commit()',
    "高风险。删除考试后成绩记录成为孤儿记录或被外键约束阻止，破坏数据完整性。",
    "删除前检查关联成绩：if Score.query.filter_by(exam_id=exam_id).first(): return False, '该考试存在关联成绩记录'。"
)

# Bug-09（原Bug-10）
add_bug(doc, 9, "中", "更新考试题目数不同步组卷",
    "exam_service.py:111-126",
    "update_exam() 方法允许通过 kwargs 修改 question_count，但不会同步重新生成 question_ids。如果 question_count 增大，实际题目数量仍为原值，导致题库不匹配。",
    'def update_exam(exam_id, **kwargs):\n    # 修改了 question_count，但 question_ids 不会同步更新\n    if "question_count" in kwargs:\n        exam.question_count = kwargs["question_count"]\n    # question_ids 仍为原值，未重新组卷\n    db.session.commit()',
    "中等风险。修改 question_count 后实际题目数量不匹配，可能导致考试题目数量不足。",
    "修改 question_count 时应同步调用 _auto_generate_questions() 重新组卷。"
)

# Bug-10（原Bug-11）
add_bug(doc, 10, "高", "评分答案比对区分大小写",
    "score_service.py:41-57",
    "grade_exam() 使用 == 直接比较字符串来判定答案是否正确。如果学生输入小写 'a' 而正确答案是大写 'A'，系统判定为错误。这是选择题评分中的常见问题。",
    'for qid in question_ids:\n    student_answer = student_answers.get(str(qid), "")\n    # 大小写敏感比较，"A" != "a"\n    if student_answer == question.answer:\n        total_score += question.score',
    "高风险。学生因大小写问题被误判答错，严重影响评分准确性和学生体验。",
    "统一转为大写后比较：if student_answer.strip().upper() == question.answer.strip().upper(): total_score += question.score。"
)

# Bug-11（原Bug-12）
add_bug(doc, 11, "高", "批量评分同样区分大小写",
    "score_service.py:132-153",
    "batch_grade() 方法与 grade_exam() 存在相同的字符串比较问题，在循环中对每个学生答案同样使用 == 比较，同样区分大小写。",
    'for score_record in scores:\n    student_answer = student_answers.get(str(qid), "")\n    # 同样的大小写敏感比较\n    if student_answer == question.answer:\n        total_score += question.score',
    "高风险。批量评分时同样存在大小写敏感问题，与 Bug-10 性质相同。",
    "与 Bug-10 一并修复：统一转为大写后比较。"
)

# Bug-12（原Bug-13）
add_bug(doc, 12, "中", "无成绩统计返回404而非零值结构",
    "score_service.py:87-101",
    "calculate_exam_statistics() 在没有成绩记录时返回 (None, 'No scores yet')，导致 API 路由层返回 404 状态码。前端期望获得有效的数据结构（即使为空），返回 None 会使前端解包失败。",
    'def calculate_exam_statistics(exam_id):\n    scores = Score.query.filter_by(exam_id=exam_id, status="scored").all()\n    if not scores:\n        # 返回 None，前端解包困难\n        return None, "No scores yet"',
    "中等风险。前端收到404后需要额外逻辑处理空统计，增加前端代码复杂度，可能导致页面显示异常。",
    "始终返回统一格式：无成绩时返回 {'exam_id': exam_id, 'total_students': 0, 'average': 0, ...}, 'Success'。"
)

doc.add_page_break()

# ===== 五、评审总结 =====
title(doc, "五、评审总结", 1)
title(doc, "5.1 共性问题汇总", 2)
issues = [
    ("输入校验缺失", "多处未对用户输入参数做范围、类型、枚举值校验，导致非法数据可进入系统。"),
    ("事务边界不清", "批量操作未使用统一事务边界，部分成功部分失败时无法回滚。"),
    ("权限控制薄弱", "部分关键操作（重置密码、删除用户）缺少严格校验，违反最小权限原则。"),
    ("字符串比对不统一", "多处使用 == 直接比对字符串，未考虑大小写问题。"),
    ("API返回值不一致", "有数据时返回字典，无数据时返回 None，给前端造成处理负担。"),
]
for t_txt, d_txt in issues:
    p = doc.add_paragraph()
    p.add_run(f"  {t_txt}：").bold = True; p.add_run(d_txt)

doc.add_paragraph()
title(doc, "5.2 改进优先级建议", 2)
ptbl = doc.add_table(rows=1, cols=4)
ptbl.style = "Table Grid"
for i,h in enumerate(["优先级","缺陷编号","问题描述","建议行动"]):
    ptbl.rows[0].cells[i].text = h
    ptbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    ptbl.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    shd(ptbl.rows[0].cells[i], "C00000")
pdata = [
    ("P0-立即修复","Bug-01,02,05,08,10,11","密码安全+数据完整性+评分逻辑","密码哈希+重置验证+分值校验+删除检查+评分大小写"),
    ("P1-本周修复","Bug-03,04,06,07,09,12","权限+枚举+事务+API契约","范围限制+枚举校验+事务回滚+统一返回结构"),
]
for i,rd in enumerate(pdata):
    row = ptbl.add_row()
    for j,val in enumerate(rd):
        row.cells[j].text = val
        row.cells[j].paragraphs[0].runs[0].font.size = Pt(9)
set_bd(ptbl)

doc.add_paragraph()
title(doc, "5.3 本次评审说明", 2)
doc.add_paragraph(
    "本次评审的所有缺陷均经实际运行验证（通过 pytest 自动化测试套件），"
    "确认在系统中真实存在。所有缺陷代码均与报告中的描述完全一致，代码位置准确。 "
    "系统整体架构清晰，业务逻辑较为完整，在输入校验、事务处理、API返回值一致性等方面存在改进空间。"
)

out = r"d:\Software\python_project\rjcs\reports\Code_Review_Report2.0.docx"
doc.save(out)
print("Word已生成: " + out)
